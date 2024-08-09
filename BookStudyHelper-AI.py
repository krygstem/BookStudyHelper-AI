# -*- coding: utf-8 -*-

import os
import time
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tqdm import tqdm
from datetime import datetime
import json
import tiktoken

# Configurar o cliente OpenAI
client = OpenAI(api_key="YOUR API")

def gerar_titulo_resumido(documento_resumido, max_palavras=3):
    texto_completo = ' '.join([p.text for p in documento_resumido.paragraphs])
    prompt = f"Baseado no seguinte texto, gere um título resumido de no máximo {max_palavras} palavras que capture o tema geral do documento:\n\n{texto_completo[:2000]}"  # Limitamos a 2000 caracteres para evitar exceder limites de tokens

    resposta = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Você é um assistente especializado em criar títulos concisos e relevantes."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=20  # Limitamos a 20 tokens para garantir um título curto
    )

    titulo = resposta.choices[0].message.content.strip()
    palavras = titulo.split()
    return ' '.join(palavras[:max_palavras])

def salvar_progresso(progresso, caminho_arquivo):
    with open(caminho_arquivo, 'w') as arquivo:
        json.dump(progresso, arquivo)

def carregar_progresso(caminho_arquivo):
    if os.path.exists(caminho_arquivo):
        with open(caminho_arquivo, 'r') as arquivo:
            return json.load(arquivo)
    return None

def perguntar_continuar():
    root = tk.Tk()
    root.withdraw()
    resposta = messagebox.askyesno("Continuar", "Deseja continuar de onde parou?")
    root.destroy()
    return resposta

def escolher_arquivos():
    root = tk.Tk()
    root.withdraw()
    caminhos_arquivos = filedialog.askopenfilenames(filetypes=[("Documentos", "*.pdf;*.doc;*.docx")])
    return sorted(caminhos_arquivos)

def criar_janela_selecao(titulo, opcoes):
    def on_select():
        nonlocal selecao
        selecao = dropdown.get()
        root.quit()

    root = tk.Tk()
    root.title(titulo)
    root.geometry("430x250")  # Aumentar a largura da janela para 250 pixels

    label = tk.Label(root, text="Escolha uma opção:")
    label.pack(pady=10)

    selecao = tk.StringVar()
    dropdown = ttk.Combobox(root, textvariable=selecao, values=opcoes, state="readonly")
    dropdown.pack(pady=10)
    dropdown.set(opcoes[0])  # Valor padrão

    botao = tk.Button(root, text="Confirmar", command=on_select)
    botao.pack(pady=10)

    root.mainloop()
    root.destroy()
    return selecao

def escolher_publico():
    opcoes = ["Crianças de 5 anos", "Crianças de 8 anos", "Pré-adolescentes de 12 anos", "Adolescentes de 15 anos", "Adultos"]
    publico = criar_janela_selecao("Escolha o público-alvo", opcoes)
    if publico == "Adultos":
        campo_profissao = pedir_profissao()
        return (publico, campo_profissao)
    return (publico, None)

def pedir_profissao():
    def on_submit():
        nonlocal profissao
        profissao = entry.get()
        if not profissao:
            profissao = "Relações Internacionais e Jogos Digitais com foco em programação"
        root.quit()

    profissao = ""
    root = tk.Tk()
    root.title("Profissão ou Área de Formação")
    root.geometry("430x350")

    label = tk.Label(root, text="Digite sua profissão ou área de formação:")
    label.pack(pady=10)

    entry = tk.Entry(root)
    entry.pack(pady=10)

    button = tk.Button(root, text="Confirmar", command=on_submit)
    button.pack(pady=10)

    root.mainloop()
    root.destroy()
    return profissao

def escolher_tokens_entrada():
    def on_submit():
        nonlocal tokens_entrada
        try:
            tokens_entrada = int(entry.get())
            if tokens_entrada < 100:
                tokens_entrada = 100
            elif tokens_entrada > 127950:
                tokens_entrada = 127950
        except ValueError:
            tokens_entrada = 100  # Valor padrão
        root.quit()

    tokens_entrada = 10000  # Valor padrão
    root = tk.Tk()
    root.title("Número de Tokens de Entrada")
    root.geometry("430x350")

    label = tk.Label(root, text="Digite o número de tokens de entrada:")
    label.pack(pady=10)

    entry = tk.Entry(root)
    entry.pack(pady=10)
    entry.insert(0, str(tokens_entrada))  # Valor padrão

    button = tk.Button(root, text="Confirmar", command=on_submit)
    button.pack(pady=10)

    root.mainloop()
    root.destroy()
    return tokens_entrada

def escolher_tokens():
    opcoes = ["1000", "2000", "4000", "8000", "15995"]
    return int(criar_janela_selecao("Escolha o número de tokens de saída", opcoes))

def escolher_tokens2():
    opcoes = ["1000", "2000", "4000", "8000", "15995"]
    return int(criar_janela_selecao("Escolha o número de tokens de saída aditivo", opcoes))

def escolher_cronologia():
    opcoes = ["Sim", "Não"]
    escolha = criar_janela_selecao("Incluir cronologia?", opcoes)
    if escolha == "Sim":
        return True
    return False

def escolher_perguntas():
    opcoes = ["Sim", "Não"]
    escolha = criar_janela_selecao("Incluir perguntas?", opcoes)
    if escolha == "Sim":
        return True
    return False

def escolher_numero_perguntas():
    def on_submit():
        nonlocal numero_perguntas
        try:
            numero_perguntas = int(entry.get())
            if numero_perguntas < 1:
                numero_perguntas = 1
            elif numero_perguntas > 50:
                numero_perguntas = 50
        except ValueError:
            numero_perguntas = 5  # Valor padrão
        root.quit()

    numero_perguntas = 5  # Valor padrão
    root = tk.Tk()
    root.title("Número de Perguntas")
    root.geometry("430x350")

    label = tk.Label(root, text="Digite o número de perguntas (1-50):")
    label.pack(pady=10)

    entry = tk.Entry(root)
    entry.pack(pady=10)
    entry.insert(0, str(numero_perguntas))  # Valor padrão

    button = tk.Button(root, text="Confirmar", command=on_submit)
    button.pack(pady=10)

    root.mainloop()
    root.destroy()
    return numero_perguntas

def ler_arquivo(caminho):
    _, extensao = os.path.splitext(caminho)
    if extensao.lower() == '.pdf':
        with open(caminho, 'rb') as arquivo:
            leitor = PdfReader(arquivo)
            return ' '.join([pagina.extract_text() for pagina in leitor.pages])
    elif extensao.lower() in ['.doc', '.docx']:
        doc = Document(caminho)
        return ' '.join([paragrafo.text for paragrafo in doc.paragraphs])

def resumir_texto2(texto, publico, max_tokens2, quantidade, incluir_cronologia, incluir_perguntas, numero_perguntas):
    quantidade = quantidade * 3  # Adicionar cálculo para multiplicar a quantidade por 2.5

    prompt = f"Adicione um dicionário de palavras-chave e explique bem cada termo, use exemplos. "

    if incluir_cronologia:
        prompt += "Adicione uma cronologia. "

    if incluir_perguntas:
        prompt += f"Adicione {numero_perguntas} perguntas de sim ou não baseadas no conteúdo, sendo que antes de cada pergunta deve estar escrito 'Responda à pergunta:' e após cada pergunta deve vir uma resposta com extensão entre 100 palavras. "

    prompt += f" Escreva tudo isso em {max_tokens2} tokens."
    prompt += f"\n\n O texto a ser trabalhado é este: {texto}"

    resposta = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.3,
        messages=[
            {"role": "system",
             "content": "Você é um assistente especializado em conteúdo para concursos públicos no Brasil. "
             "Vôcê não escreve qualquer informação sobre seus procedimentos internos, limitações ou porque fará algo. "
             "Você evita o uso de siglas ao máximo. " 
             "Você se atêm totalmente ao texto original."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=max_tokens2
    )

    return resposta.choices[0].message.content

def resumir_texto(texto, publico, max_tokens, quantidade, incluir_cronologia, incluir_perguntas, numero_perguntas):
    quantidade = quantidade * 3  # Adicionar cálculo para multiplicar a quantidade por 2.5

    prompt = f"Adapte o seguinte texto para que ocupe {quantidade} caracteres. "
    prompt += f"O texto deve ser escrito para {publico}. "
    prompt += "Preserve todos os conceitos, números, nomes, locais e datas. "
    prompt += "Transforme todas as perguntas do texto original em afirmações. "
    prompt += f" Escreva tudo isso em {max_tokens} tokens."
    prompt += f"\n\n O texto a ser trabalhado é este: {texto}"

    resposta = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.1,
        messages=[
            {"role": "system",
             "content": "Você é um assistente especializado em conteúdo para concursos públicos no Brasil. "
             "Vôcê não escreve qualquer informação sobre seus procedimentos internos, limitações ou porque fará algo. "
             "Você evita o uso de siglas ao máximo. " 
             "Você se atêm totalmente ao texto original."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=max_tokens
    )

    return resposta.choices[0].message.content

def calcular_custos(arquivos, max_tokens, max_tokens2):
    num_paginas = 0

    for arquivo in arquivos:
        if arquivo.lower().endswith('.pdf'):
            with open(arquivo, 'rb') as f:
                leitor_pdf = PdfReader(f)
                num_paginas += len(leitor_pdf.pages)  # Use len(reader.pages)

        elif arquivo.lower().endswith(('.doc', '.docx')):
            doc = Document(arquivo)
            num_paginas += len(doc.paragraphs) // 2  # Estimativa para documentos do Word

    total_tokens = 1000 * num_paginas
    custo = (total_tokens / 1_000_000) * 0.150 + (16000 / 1_000_000) * 0.600
    tempo_estimado = total_tokens / 1000 * 5 + (len(arquivos) * 5 * (total_tokens / max_tokens + max_tokens2))  # Tempo de trabalho mais pausas

    return total_tokens, custo, tempo_estimado

def substituir_caracteres(documento):
    for paragrafo in documento.paragraphs:
        texto_original = paragrafo.text
        texto_modificado = texto_original.replace('***', '*').replace('###', '#')
        if texto_original != texto_modificado:
            paragrafo.text = texto_modificado

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto_original = paragrafo.text
                    texto_modificado = texto_original.replace('*', '...').replace('#', '...')
                    if texto_original != texto_modificado:
                        paragrafo.text = texto_modificado

def main():
    progresso_path = "progresso.json"
    progresso = carregar_progresso(progresso_path)

    if progresso and perguntar_continuar():
        caminhos_arquivos = progresso['caminhos_arquivos']
        publico = progresso['publico']
        profissao = progresso['profissao']
        max_tokens = progresso['max_tokens']
        max_tokens2 = progresso['max_tokens2']
        tokens_entrada = progresso['tokens_entrada']
        index_arquivo = progresso['index_arquivo']
        index_chunk = progresso['index_chunk']
        documento_resumido_path = progresso['documento_resumido_path']
        documento_resumido = Document(documento_resumido_path)
        caminho_saida = documento_resumido_path
        incluir_cronologia = progresso['incluir_cronologia']
        incluir_perguntas = progresso['incluir_perguntas']
        numero_perguntas = progresso['numero_perguntas']
    else:
        caminhos_arquivos = escolher_arquivos()
        if not caminhos_arquivos:
            print("Nenhum arquivo selecionado. Encerrando o programa.")
            return

        publico, profissao = escolher_publico()
        if not publico:
            print("Nenhum público-alvo selecionado. Encerrando o programa.")
            return

        max_tokens = escolher_tokens()
        if not max_tokens:
            print("Nenhum limite de tokens selecionado. Encerrando o programa.")
            return

        max_tokens2 = escolher_tokens2()
        if not max_tokens2:
            print("Nenhum limite de tokens selecionado. Encerrando o programa.")
            return

        tokens_entrada = escolher_tokens_entrada()
        if not tokens_entrada:
            print("Nenhum número de tokens de entrada selecionado. Encerrando o programa.")
            return

        documento_resumido = Document()
        index_arquivo = 0
        index_chunk = 0
        pasta_origem = os.path.dirname(caminhos_arquivos[0])
        caminho_saida = os.path.join(pasta_origem, "documento_temporario.docx")

        # Menu inicial
        incluir_cronologia = escolher_cronologia()
        incluir_perguntas = escolher_perguntas()
        if incluir_perguntas:
            numero_perguntas = escolher_numero_perguntas()
        else:
            numero_perguntas = 0

        # Adicionar informações no início do documento resumido
        data_hora_inicio = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        nomes_arquivos = ", ".join([os.path.basename(arquivo) for arquivo in caminhos_arquivos])
        documento_resumido.add_paragraph(f"Data e hora de início: {data_hora_inicio}")
        documento_resumido.add_paragraph(f"Arquivos trabalhados: {nomes_arquivos}")
        documento_resumido.add_paragraph(f"Público-alvo: {publico}")
        documento_resumido.add_paragraph(f"Profissão: {profissao}")
        documento_resumido.add_paragraph(f"Máximo de tokens de saída do resumo: {max_tokens}")
        documento_resumido.add_paragraph(f"Máximo de tokens de saída aditivos: {max_tokens2}")
        documento_resumido.add_paragraph(f"Tokens de entrada: {tokens_entrada}")
        documento_resumido.add_paragraph(f"Incluir cronologia: {incluir_cronologia}")
        documento_resumido.add_paragraph(f"Incluir perguntas: {incluir_perguntas}")
        documento_resumido.add_paragraph(f"Número de perguntas: {numero_perguntas}")
        documento_resumido.add_paragraph("\n")

    # Calcular quantidade
    quantidade = max_tokens * 0.7

    total_tokens, custo, tempo_estimado = calcular_custos(caminhos_arquivos, max_tokens, max_tokens2)

    info_msg = (f"1 - Total de tokens dos arquivos: {total_tokens}\n"
                f"2 - Custo simulado da tarefa final em dólares: ${custo:.2f}\n"
                f"3 - Tempo aproximado da tarefa: {tempo_estimado / 60:.2f} minutos\n\n"
                "Pressione OK para continuar.")
    messagebox.showinfo("Informações de Tarefa", info_msg)

    for i in tqdm(range(index_arquivo, len(caminhos_arquivos)), desc="Processando arquivos"):
        caminho_arquivo = caminhos_arquivos[i]
        texto_completo = ler_arquivo(caminho_arquivo)

        if i == index_arquivo:
            documento_resumido.add_heading(f"Resumo do arquivo: {os.path.basename(caminho_arquivo)}", level=1)

        chunk_size = tokens_entrada
        chunks = [texto_completo[j:j + chunk_size] for j in range(0, len(texto_completo), chunk_size)]

        for j in range(index_chunk, len(chunks)):
            chunk = chunks[j]
            print(f"Resumindo parte {j + 1} do arquivo {os.path.basename(caminho_arquivo)}...")
            #resumo = resumir_texto(chunk, publico, max_tokens, quantidade, incluir_cronologia, incluir_perguntas, numero_perguntas)
            #documento_resumido.add_paragraph(resumo)
            resumo1 = resumir_texto(chunk, publico, max_tokens, quantidade, incluir_cronologia, incluir_perguntas,
                                    numero_perguntas)
            resumo2 = resumir_texto2(chunk, publico, max_tokens2, quantidade, incluir_cronologia, incluir_perguntas,
                                     numero_perguntas)

            documento_resumido.add_paragraph(resumo1)
            documento_resumido.add_paragraph("")  # Adiciona um parágrafo vazio
            documento_resumido.add_paragraph(resumo2)
            time.sleep(0.1)

            # Salvar progresso
            progresso = {
                'caminhos_arquivos': caminhos_arquivos,
                'publico': publico,
                'profissao': profissao,
                'max_tokens': max_tokens,
                'max_tokens2': max_tokens2,
                'tokens_entrada': tokens_entrada,
                'index_arquivo': i,
                'index_chunk': j + 1,
                'documento_resumido_path': caminho_saida,
                'incluir_cronologia': incluir_cronologia,
                'incluir_perguntas': incluir_perguntas,
                'numero_perguntas': numero_perguntas
            }
            documento_resumido.save(caminho_saida)
            salvar_progresso(progresso, progresso_path)

        index_chunk = 0  # Resetar o índice do chunk para o próximo arquivo

    # Substituir asteriscos e cerquilhas por reticências
    substituir_caracteres(documento_resumido)

    # Gerar título resumido
    titulo_resumido = gerar_titulo_resumido(documento_resumido)

    # Criar nome do arquivo final com o título resumido
    now = datetime.now()
    filename = f"{titulo_resumido} {now.strftime('%H-%M')} {now.strftime('%d-%m-%Y')}.docx"
    caminho_saida_final = os.path.join(os.path.dirname(caminho_saida), filename)

    documento_resumido.save(caminho_saida_final)
    messagebox.showinfo("Concluído", f"Resumo salvo em: {caminho_saida_final}")

    # Remover o arquivo de progresso e o arquivo temporário após a conclusão
    if os.path.exists(progresso_path):
        os.remove(progresso_path)
    if os.path.exists(caminho_saida) and caminho_saida != caminho_saida_final:
        os.remove(caminho_saida)

if __name__ == "__main__":
    main()
